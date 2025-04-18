from fastapi import FastAPI, HTTPException, Depends, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
from sqlalchemy import Column, Integer, String, create_engine
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
import torch
import re
import string
import secrets
import os
import io
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document as DocxDocument
import docx

# ---------- Transformer Model ----------
class Vocabulary:
    def __init__(self):
        self.word2idx = {}
        self.idx2word = {}
        self.size = 0

class TransformerTranslator(torch.nn.Module):
    def __init__(self, src_vocab_size, tgt_vocab_size,
                 embed_dim=256, num_heads=8, ff_dim=2048,
                 num_layers=3, dropout=0.1, max_len=500):
        super().__init__()
        self.embed_dim = embed_dim
        self.num_heads = num_heads
        self.ff_dim = ff_dim
        self.num_layers = num_layers
        self.dropout = dropout
        self.max_len = max_len

        self.encoder_embed = torch.nn.Embedding(src_vocab_size, embed_dim)
        self.decoder_embed = torch.nn.Embedding(tgt_vocab_size, embed_dim)
        self.pos_encoder = torch.nn.Embedding(max_len, embed_dim)

        self.transformer = torch.nn.Transformer(
            d_model=embed_dim,
            nhead=num_heads,
            num_encoder_layers=num_layers,
            num_decoder_layers=num_layers,
            dim_feedforward=ff_dim,
            dropout=dropout
        )
        self.fc_out = torch.nn.Linear(embed_dim, tgt_vocab_size)

    def forward(self, src, tgt):
        src_pos = torch.arange(0, src.size(1), device=src.device).unsqueeze(0)
        tgt_pos = torch.arange(0, tgt.size(1), device=tgt.device).unsqueeze(0)

        src_emb = self.encoder_embed(src) + self.pos_encoder(src_pos)
        tgt_emb = self.decoder_embed(tgt) + self.pos_encoder(tgt_pos)

        tgt_mask = torch.nn.Transformer.generate_square_subsequent_mask(tgt.size(1)).to(src.device)

        output = self.transformer(
            src_emb.permute(1, 0, 2),
            tgt_emb.permute(1, 0, 2),
            tgt_mask=tgt_mask
        )
        return self.fc_out(output.permute(1, 0, 2))

# ---------- Database ----------
DATABASE_URL = 'postgresql://postgres:postgres@localhost/Translator'
ADMIN_TOKEN = os.getenv("ADMIN_TOKEN", "supersecretadmintoken")

Base = declarative_base()

class User(Base):
    __tablename__ = 'users'
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String(50), unique=True, index=True)
    password = Column(String(50))
    token = Column(String(32), unique=True, index=True)
    attempts_left = Column(Integer, default=10)

engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(bind=engine)

def create_db():
    Base.metadata.create_all(bind=engine)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ---------- App & Model ----------
async def lifespan(app: FastAPI):
    global model, src_vocab, tgt_vocab, device
    create_db()
    try:
        device = 'cuda' if torch.cuda.is_available() else 'cpu'
        model, src_vocab, tgt_vocab, _ = load_model('data/best_model.pth', device)
        print("Model loaded successfully")
    except Exception as e:
        print(f"Error loading model: {str(e)}")
        raise
    yield
    if torch.cuda.is_available():
        torch.cuda.empty_cache()

app = FastAPI(
    title="Translation API with Auth",
    description="API for text translation with authentication",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def load_model(model_path, device='cpu'):
    checkpoint = torch.load(model_path, map_location=device)
    src_vocab = Vocabulary()
    src_vocab.__dict__ = checkpoint['src_vocab']
    tgt_vocab = Vocabulary()
    tgt_vocab.__dict__ = checkpoint['tgt_vocab']
    params = checkpoint['model_params']
    model = TransformerTranslator(
        src_vocab_size=src_vocab.size,
        tgt_vocab_size=tgt_vocab.size,
        **params
    )
    model.load_state_dict(checkpoint['model_state_dict'])
    model.eval()
    return model, src_vocab, tgt_vocab, params

def preprocess(text, vocab, max_len=500):
    text = re.sub(f'[{re.escape(string.punctuation)}]', '', text.lower())
    words = text.split()[:max_len]
    return [vocab.word2idx.get(word, 1) for word in words]

def translate_text(model, src_vocab, tgt_vocab, text, device='cpu', max_len=500):
    model.to(device)
    src = torch.LongTensor(preprocess(text, src_vocab, max_len)).unsqueeze(0).to(device)
    output = [tgt_vocab.word2idx['[start]']]
    for _ in range(max_len):
        tgt = torch.LongTensor(output).unsqueeze(0).to(device)
        with torch.no_grad():
            pred = model(src, tgt)
        next_idx = pred[0, -1].argmax().item()
        if next_idx == tgt_vocab.word2idx['[end]']:
            break
        output.append(next_idx)
    translated = ' '.join([tgt_vocab.idx2word.get(idx, '<unk>') for idx in output])
    return translated.replace('[start]', '').replace('[end]', '').strip()

# ---------- Document Processing ----------
ALLOWED_FILE_TYPES = {'pdf', 'docx', 'txt'}

def process_file(file: UploadFile):
    content = file.file.read()
    file_type = file.filename.split('.')[-1].lower()
    
    if file_type == 'pdf':
        text = extract_pdf_text(io.BytesIO(content))
    elif file_type == 'docx':
        doc = DocxDocument(io.BytesIO(content))
        text = '\n'.join([para.text for para in doc.paragraphs])
    elif file_type == 'txt':
        text = content.decode('utf-8')
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    return text, file_type

def recreate_docx(translated_text, original_doc):
    new_doc = DocxDocument()
    for paragraph in original_doc.paragraphs:
        new_p = new_doc.add_paragraph()
        for run in paragraph.runs:
            new_run = new_p.add_run(translate_text(model, src_vocab, tgt_vocab, run.text))
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
    return new_doc

# ---------- Pydantic Models ----------
class VerifyTokenRequest(BaseModel):
    token: str

class TranslationRequest(BaseModel):
    token: str
    text: str
    max_length: Optional[int] = 500

class TranslationResponse(BaseModel):
    translation: str
    model_device: str
    attempts_left: int

class FileTranslationResponse(BaseModel):
    translated_file: bytes
    filename: str

class UserCreate(BaseModel):
    username: str
    password: str

class TokenInfo(BaseModel):
    username: str
    token: str
    attempts_left: int

class RefillRequest(BaseModel):
    admin_token: str
    user_token: str
    attempts: int

class RefillResponse(BaseModel):
    message: str
    new_attempts: int

# ---------- Endpoints ----------
@app.post("/register", response_model=TokenInfo)
def register(user: UserCreate, db: Session = Depends(get_db)):
    if db.query(User).filter(User.username == user.username).first():
        raise HTTPException(status_code=400, detail="Username already exists")
    
    token = secrets.token_hex(16)
    db_user = User(
        username=user.username,
        password=user.password,
        token=token
    )
    
    db.add(db_user)
    try:
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail="Database error")
    
    db.refresh(db_user)
    return {
        "username": db_user.username,
        "token": db_user.token,
        "attempts_left": db_user.attempts_left
    }

@app.post("/login", response_model=TokenInfo)
def login(user_data: UserCreate, db: Session = Depends(get_db)):
    user = db.query(User).filter(
        User.username == user_data.username,
        User.password == user_data.password
    ).first()
    
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    return {
        "username": user.username,
        "token": user.token,
        "attempts_left": user.attempts_left
    }

@app.post("/verify-token", response_model=TokenInfo)
def verify_token(request: VerifyTokenRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.token == request.token).first()
    if not user:
        raise HTTPException(status_code=404, detail="Invalid token")
    return {
        "username": user.username,
        "token": user.token,
        "attempts_left": user.attempts_left
    }

@app.get("/token/{username}", response_model=TokenInfo)
def get_token(username: str, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return {
        "username": user.username,
        "token": user.token,
        "attempts_left": user.attempts_left
    }

@app.post("/translate", response_model=TranslationResponse)
def api_translate(
    request: TranslationRequest,
    db: Session = Depends(get_db)
):
    user = db.query(User).filter(User.token == request.token).first()
    if not user:
        raise HTTPException(status_code=401, detail="Invalid token")
    if user.attempts_left <= 0:
        raise HTTPException(status_code=403, detail="No attempts left")
    if len(request.text.split()) > 500:
        raise HTTPException(status_code=400, detail="Text exceeds 500 words limit")

    try:
        translation = translate_text(
            model,
            src_vocab,
            tgt_vocab,
            request.text,
            device=device,
            max_len=request.max_length
        )
        user.attempts_left -= 1
        db.commit()
        return {
            "translation": translation,
            "model_device": device,
            "attempts_left": user.attempts_left
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/translate-file", response_model=FileTranslationResponse)
async def translate_file(
    token: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    user = db.query(User).filter(User.token == token).first()
    if not user:
        raise HTTPException(status_code=401, detail="Invalid token")
    if user.attempts_left <= 0:
        raise HTTPException(status_code=403, detail="No attempts left")

    try:
        original_text, file_type = process_file(file)
        translated_text = translate_text(model, src_vocab, tgt_vocab, original_text)
        
        output = io.BytesIO()
        if file_type == 'docx':
            original_doc = DocxDocument(io.BytesIO(await file.read()))
            new_doc = recreate_docx(translated_text, original_doc)
            new_doc.save(output)
        else:
            output.write(translated_text.encode('utf-8'))
        
        user.attempts_left -= 1
        db.commit()
        
        return {
            "translated_file": output.getvalue(),
            "filename": f"translated_{file.filename}"
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/refill-attempts", response_model=RefillResponse)
def refill_attempts(
    request: RefillRequest,
    db: Session = Depends(get_db)
):
    if request.admin_token != ADMIN_TOKEN:
        raise HTTPException(status_code=401, detail="Invalid admin token")
    if request.attempts <= 0:
        raise HTTPException(status_code=400, detail="Attempts must be positive")
    
    user = db.query(User).filter(User.token == request.user_token).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    user.attempts_left += request.attempts
    try:
        db.commit()
        db.refresh(user)
        return {
            "message": "Attempts refilled successfully",
            "new_attempts": user.attempts_left
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("serv:app", host="0.0.0.0", port=8000, reload=True)
