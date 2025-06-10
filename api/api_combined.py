import sys
import os
import traceback
import re
import string
import secrets
import io
import time
import threading
import logging
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed

import torch
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, BackgroundTasks, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
from sqlalchemy import Column, Integer, String, create_engine
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session

from googletrans import Translator
from docx import Document as DocxDocument
import docx
from pdfminer.high_level import extract_text as extract_pdf_text

from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml import parse_xml

import fitz  
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader

# ---------- Трансформер Модель ----------
class Vocabulary:
    def __init__(self):
        self.word2idx = {}
        self.idx2word = {}
        self.size = 0

class TransformerTranslator(torch.nn.Module):
    def __init__(self, src_vocab_size, tgt_vocab_size,
                 embed_dim=256, num_heads=8, ff_dim=2048,
                 num_layers=3, dropout=0.1, max_len=500):
        super().__init__()
        self.embed_dim = embed_dim
        self.num_heads = num_heads
        self.ff_dim = ff_dim
        self.num_layers = num_layers
        self.dropout = dropout
        self.max_len = max_len

        self.encoder_embed = torch.nn.Embedding(src_vocab_size, embed_dim)
        self.decoder_embed = torch.nn.Embedding(tgt_vocab_size, embed_dim)
        self.pos_encoder = torch.nn.Embedding(max_len, embed_dim)

        self.transformer = torch.nn.Transformer(
            d_model=embed_dim,
            nhead=num_heads,
            num_encoder_layers=num_layers,
            num_decoder_layers=num_layers,
            dim_feedforward=ff_dim,
            dropout=dropout
        )
        self.fc_out = torch.nn.Linear(embed_dim, tgt_vocab_size)

    def forward(self, src, tgt):
        src_pos = torch.arange(0, src.size(1), device=src.device).unsqueeze(0)
        tgt_pos = torch.arange(0, tgt.size(1), device=tgt.device).unsqueeze(0)

        src_emb = self.encoder_embed(src) + self.pos_encoder(src_pos)
        tgt_emb = self.decoder_embed(tgt) + self.pos_encoder(tgt_pos)

        tgt_mask = torch.nn.Transformer.generate_square_subsequent_mask(tgt.size(1)).to(src.device)

        output = self.transformer(
            src_emb.permute(1, 0, 2),
            tgt_emb.permute(1, 0, 2),
            tgt_mask=tgt_mask
        )
        return self.fc_out(output.permute(1, 0, 2))

def load_model(model_path, device='cpu'):
    checkpoint = torch.load(model_path, map_location=device)
    src_vocab = Vocabulary()
    src_vocab.__dict__ = checkpoint['src_vocab']
    tgt_vocab = Vocabulary()
    tgt_vocab.__dict__ = checkpoint['tgt_vocab']
    params = checkpoint['model_params']
    model = TransformerTranslator(
        src_vocab_size=src_vocab.size,
        tgt_vocab_size=tgt_vocab.size,
        **params
    )
    model.load_state_dict(checkpoint['model_state_dict'])
    model.eval()
    return model, src_vocab, tgt_vocab, params

def preprocess(text, vocab, max_len=500):
    text = re.sub(f'[{re.escape(string.punctuation)}]', '', text.lower())
    words = text.split()[:max_len]
    return [vocab.word2idx.get(word, 1) for word in words]

def translate_text(model, src_vocab, tgt_vocab, text, device='cpu', max_len=500):
    model.to(device)
    src = torch.LongTensor(preprocess(text, src_vocab, max_len)).unsqueeze(0).to(device)
    output = [tgt_vocab.word2idx['[start]']]
    for _ in range(max_len):
        tgt = torch.LongTensor(output).unsqueeze(0).to(device)
        with torch.no_grad():
            pred = model(src, tgt)
        next_idx = pred[0, -1].argmax().item()
        if next_idx == tgt_vocab.word2idx['[end]']:
            break
        output.append(next_idx)
    translated = ' '.join([tgt_vocab.idx2word.get(idx, '<unk>') for idx in output])
    return translated.replace('[start]', '').replace('[end]', '').strip()

# ---------- Безопасный перевод с Google Translate ----------
translator = Translator()
translator.raise_Exception = False  # Чтобы избежать ошибки AttributeError

def safe_translate_chunk(chunk):
    try:
        result = translator.translate(chunk, dest='en')
        if hasattr(result, 'extra') and 'parsed' in result.extra:
            parsed = result.extra['parsed']
            if isinstance(parsed, list):
                texts = [part.text for part in parsed if part.text is not None and isinstance(part.text, str)]
                return ' '.join(texts)
        if result.text is not None and isinstance(result.text, str):
            return result.text
        else:
            return chunk
    except Exception as e:
        print(f"Translation error: {str(e)}")
        traceback.print_exc()
        return chunk

def safe_translate(text, max_chunk_size=500, max_workers=5):
    if not text or text.isspace():
        return ""
    chunks = [text[i:i+max_chunk_size] for i in range(0, len(text), max_chunk_size)]
    translated_chunks = [""] * len(chunks)

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(safe_translate_chunk, chunk): idx for idx, chunk in enumerate(chunks)}
        for future in as_completed(futures):
            idx = futures[future]
            try:
                translated_chunks[idx] = future.result()
            except Exception as e:
                print(f"Translation error in future: {str(e)}")
                translated_chunks[idx] = chunks[idx]

    return "".join(translated_chunks)

# ---------- Копирование форматирования параграфов и текстовых фрагментов ----------
def copy_paragraph_formatting(src, dst):
    if hasattr(src.style, 'name'):
        dst.style = src.style.name
    else:
        dst.style = src.style

    dst.alignment = src.alignment

    dst_format = dst.paragraph_format
    src_format = src.paragraph_format
    dst_format.left_indent = src_format.left_indent
    dst_format.right_indent = src_format.right_indent
    dst_format.first_line_indent = src_format.first_line_indent
    dst_format.space_before = src_format.space_before
    dst_format.space_after = src_format.space_after
    dst_format.line_spacing = src_format.line_spacing
    dst_format.line_spacing_rule = src_format.line_spacing_rule

    if src._p.pPr is not None and src._p.pPr.numPr is not None:
        num_xml = src._p.pPr.numPr.xml
        dst_num_pr = parse_xml(num_xml)
        pPr = dst._p.get_or_add_pPr()
        for child in pPr.findall(qn('w:numPr')):
            pPr.remove(child)
        pPr.append(dst_num_pr)

def copy_run_formatting(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size

    if src_run.font.color and src_run.font.color.rgb:
        if isinstance(src_run.font.color.rgb, str):
            dst_run.font.color.rgb = RGBColor.from_string(src_run.font.color.rgb)
        else:
            dst_run.font.color.rgb = src_run.font.color.rgb

    if dst_run._element.rPr is not None:
        rPr = dst_run._element.rPr
        rFonts = rPr.rFonts
        if rFonts is not None and src_run.font.name:
            rFonts.set(qn('w:eastAsia'), src_run.font.name)

def copy_picture(src_run, dst_doc, parent_para=None):
    try:
        blip = src_run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        if not blip:
            print("Не найден элемент blip в run")
            return None
        blip_rId = blip[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        image_part = src_run.part.related_parts[blip_rId]
        image_bytes = image_part.blob
        image_stream = BytesIO(image_bytes)
        if parent_para:
            run = parent_para.add_run()
            run.add_picture(image_stream)
            return run
        else:
            image = dst_doc.add_picture(image_stream)
            return image
    except Exception as e:
        print(f"Ошибка копирования изображения: {e}")
        return None

# ---------- Перевод параграфа и таблицы ----------

def translate_paragraph(src_para, dst_doc):
    new_para = dst_doc.add_paragraph()
    copy_paragraph_formatting(src_para, new_para)

    for run in src_para.runs:
        drawing_elements = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        if drawing_elements:
            copy_picture(run, dst_doc, parent_para=new_para)
        else:
            text_to_translate = run.text if run.text is not None else ""
            translated_text = safe_translate(text_to_translate)
            new_run = new_para.add_run(translated_text)
            copy_run_formatting(run, new_run)

    return new_para

def translate_table(table, dst_doc):
    new_table = dst_doc.add_table(rows=0, cols=len(table.columns))
    new_table.style = table.style

    for row in table.rows:
        new_row = new_table.add_row()
        for cell_idx, cell in enumerate(row.cells):
            new_cell = new_row.cells[cell_idx]
            new_cell._element.clear_content()
            for para in cell.paragraphs:
                translated_para = translate_paragraph(para, dst_doc)
                new_cell._element.append(translated_para._element)

    return new_table

# ---------- Перевод DOCX ----------

def translate_docx(input_path, output_path, progress_callback=None):
    try:
        src_doc = DocxDocument(input_path)
        dst_doc = DocxDocument()

        body_elements = list(src_doc.element.body)
        total_elements = len(body_elements)
        current_step = 0

        for element in body_elements:
            if element.tag.endswith('}p'):
                para = src_doc.paragraphs[0]
                for p in src_doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                translate_paragraph(para, dst_doc)
            elif element.tag.endswith('}tbl'):
                tbl = src_doc.tables[0]
                for t in src_doc.tables:
                    if t._element == element:
                        tbl = t
                        break
                translate_table(tbl, dst_doc)
            elif element.tag.endswith('}drawing'):
                pass
            current_step += 1
            if progress_callback:
                progress_callback(int(current_step / total_elements * 100))

        dst_doc.save(output_path)
        return True
    except Exception as e:
        print(f"DOCX Translation Error: {str(e)}")
        return False

# ---------- Перевод PDF ----------

def translate_pdf(input_path, output_path, progress_callback=None):
    try:
        doc = fitz.open(input_path)
        new_pdf = fitz.open()

        total_pages = len(doc)
        for page_num in range(total_pages):
            page = doc.load_page(page_num)
            blocks = page.get_text("dict")["blocks"]

            new_page = new_pdf.new_page(width=page.rect.width, height=page.rect.height)

            for b in blocks:
                if b['type'] == 0:
                    for line in b["lines"]:
                        line_text = " ".join([span["text"] for span in line["spans"]])
                        translated_text = safe_translate(line_text)
                        new_page.insert_text((b["bbox"][0], b["bbox"][1]), translated_text,
                                             fontsize=12, fontname="helv")
                elif b['type'] == 1:
                    pass
            if progress_callback:
                progress_callback(int((page_num + 1) / total_pages * 100))

        new_pdf.save(output_path)
        new_pdf.close()
        doc.close()
        return True
    except Exception as e:
        print(f"PDF Translation Error: {str(e)}")
        return False

# ---------- Перевод TXT ----------
def translate_txt(input_path, output_path):
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        translated = safe_translate(content)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(translated)

        return True
    except Exception as e:
        print(f"TXT Translation Error: {str(e)}")
        return False

# ---------- База данных ----------
DATABASE_URL = 'postgresql://postgres:postgres@localhost/Translator'
ADMIN_TOKEN = os.getenv("ADMIN_TOKEN", "supersecretadmintoken")

Base = declarative_base()

class User(Base):
    __tablename__ = 'users'
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String(50), unique=True, index=True)
    password = Column(String(50))
    token = Column(String(32), unique=True, index=True)
    attempts_left = Column(Integer, default=10)

engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(bind=engine)

from sqlalchemy.orm import declarative_base

Base = declarative_base()

def create_db():
    Base.metadata.create_all(bind=engine)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ---------- Приложение и модель ----------

async def lifespan(app: FastAPI):
    global model, src_vocab, tgt_vocab, device, translator
    create_db()
    try:
        device = 'cuda' if torch.cuda.is_available() else 'cpu'
        model, src_vocab, tgt_vocab, _ = load_model('data/best_model.pth', device)
        print("Model loaded successfully")
        translator = Translator()
    except Exception as e:
        print(f"Error loading model: {str(e)}")
        raise
    yield
    if torch.cuda.is_available():
        torch.cuda.empty_cache()

app = FastAPI(
    title="Translation API with Auth",
    description="API для перевода текста с аутентификацией",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- Обработка файлов ----------
ALLOWED_FILE_TYPES = {'pdf', 'docx', 'txt'}

def process_docx(input_path):
    try:
        src_doc = DocxDocument(input_path)
        dst_doc = DocxDocument()
        
        for para in src_doc.paragraphs:
            translate_paragraph(para, dst_doc)
        
        for table in src_doc.tables:
            translate_table(table, dst_doc)
        
        output = io.BytesIO()
        dst_doc.save(output)
        output.seek(0)
        return output.read()
    except Exception as e:
        print(f"DOCX Translation Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

def process_file(file: UploadFile):
    content = file.file.read()
    file_type = file.filename.split('.')[-1].lower()
    
    if file_type == 'docx':
        return process_docx(io.BytesIO(content)), file_type
    elif file_type == 'pdf':
        text = extract_pdf_text(io.BytesIO(content))
        translated_text = translate_text(model, src_vocab, tgt_vocab, text)
        return translated_text.encode('utf-8'), file_type
    elif file_type == 'txt':
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            text = content.decode('latin-1')
        translated_text = translate_text(model, src_vocab, tgt_vocab, text)
        return translated_text.encode('utf-8'), file_type
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

# ---------- Pydantic модели ----------
class VerifyTokenRequest(BaseModel):
    token: str

class TranslationRequest(BaseModel):
    token: str
    text: str
    max_length: Optional[int] = 10
    method: Optional[str] = "model"

class TranslationResponse(BaseModel):
    translation: str
    model_device: str
    attempts_left: int

class FileTranslationResponse(BaseModel):
    translated_file: bytes
    filename: str

class UserCreate(BaseModel):
    username: str
    password: str

class TokenInfo(BaseModel):
    username: str
    token: str
    attempts_left: int

class RefillRequest(BaseModel):
    admin_token: str
    user_token: str
    attempts: int

class RefillResponse(BaseModel):
    message: str
    new_attempts: int

# ---------- Эндпоинты ----------
import hashlib
import base64

def hash_password(password: str) -> str:
    # Используем sha256 и кодируем в base64 для сокращения длины
    sha256_hash = hashlib.sha256(password.encode('utf-8')).digest()
    b64_hash = base64.b64encode(sha256_hash).decode('utf-8')
    # Обрезаем до 50 символов, чтобы уместиться в поле базы данных
    return b64_hash[:50]

@app.post("/register", response_model=TokenInfo)
def register(user: UserCreate, db: Session = Depends(get_db)):
    if db.query(User).filter(User.username == user.username).first():
        raise HTTPException(status_code=400, detail="Имя пользователя уже существует")
    
    token = secrets.token_hex(16)
    hashed_password = hash_password(user.password)
    db_user = User(
        username=user.username,
        password=hashed_password,
        token=token,
        attempts_left=10
    )
    
    db.add(db_user)
    try:
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Ошибка базы данных: {str(e)}")
    
    db.refresh(db_user)
    return {
        "username": db_user.username,
        "token": db_user.token,
        "attempts_left": db_user.attempts_left
    }

def verify_password(plain_password: str, hashed_password: str) -> bool:
    import hashlib
    import base64
    sha256_hash = hashlib.sha256(plain_password.encode('utf-8')).digest()
    b64_hash = base64.b64encode(sha256_hash).decode('utf-8')
    return b64_hash[:50] == hashed_password

@app.post("/login", response_model=TokenInfo)
def login(user_data: UserCreate, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == user_data.username).first()
    if not user or not verify_password(user_data.password, user.password):
        raise HTTPException(status_code=401, detail="Неверные учетные данные")
    
    return {
        "username": user.username,
        "token": user.token,
        "attempts_left": user.attempts_left
    }

@app.post("/verify-token", response_model=TokenInfo)
def verify_token(request: VerifyTokenRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.token == request.token).first()
    if not user:
        raise HTTPException(status_code=404, detail="Неверный токен")
    return {
        "username": user.username,
        "token": user.token,
        "attempts_left": user.attempts_left
    }

@app.get("/token/{username}", response_model=TokenInfo)
def get_token(username: str, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    return {
        "username": user.username,
        "token": user.token,
        "attempts_left": user.attempts_left
    }

@app.post("/translate", response_model=TranslationResponse)
def api_translate(
    request: TranslationRequest,
    db: Session = Depends(get_db)
):
    user = db.query(User).filter(User.token == request.token).first()
    if not user:
        raise HTTPException(status_code=401, detail="Неверный токен")
    if user.attempts_left <= 0:
        raise HTTPException(status_code=403, detail="Попытки закончились")
    if len(request.text.split()) > 10:
        raise HTTPException(status_code=400, detail="Текст превышает лимит в 500 слов")

    try:
        if request.method == "google":
            translation = safe_translate(request.text)
            model_device = "googletrans"
        else:
            translation = translate_text(
                model,
                src_vocab,
                tgt_vocab,
                request.text,
                device=device,
                max_len=request.max_length
            )
            model_device = device

        user.attempts_left -= 1
        db.commit()
        return {
            "translation": translation,
            "model_device": model_device,
            "attempts_left": user.attempts_left
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

progress_dict = {}

def translate_file_background(task_id, file_bytes, filename, method):
    try:
        logging.info(f"Начало перевода для task_id={task_id}, filename={filename}, method={method}")
        if method == "google":
            file_stream = BytesIO(file_bytes)
            file_stream.name = filename

            def progress_callback(value):
                progress_dict[task_id] = value

            ext = filename.split('.')[-1].lower()
            if ext == 'docx':
                output_path = f"temp_{task_id}.docx"
                success = translate_docx(file_stream, output_path, progress_callback=progress_callback)
                if success:
                    with open(output_path, 'rb') as f:
                        translated_content = f.read()
                    progress_dict[task_id] = 100
                    progress_dict[f"{task_id}_result"] = (translated_content, ext)
                    logging.info(f"DOCX перевод успешен для task_id={task_id}, размер файла={len(translated_content)} байт, ext={ext}")
                    os.remove(output_path)
                    logging.info(f"Временный файл удален для task_id={task_id}")
                else:
                    progress_dict[f"{task_id}_error"] = "Ошибка перевода DOCX"
                    logging.error(f"Ошибка перевода DOCX для task_id={task_id}")
            elif ext == 'pdf':
                output_path = f"temp_{task_id}.pdf"
                success = translate_pdf(file_stream, output_path, progress_callback=progress_callback)
                if success:
                    with open(output_path, 'rb') as f:
                        translated_content = f.read()
                    progress_dict[task_id] = 100
                    progress_dict[f"{task_id}_result"] = (translated_content, ext)
                    os.remove(output_path)
                    logging.info(f"PDF перевод успешен для task_id={task_id}")
                else:
                    progress_dict[f"{task_id}_error"] = "Ошибка перевода PDF"
                    logging.error(f"Ошибка перевода PDF для task_id={task_id}")
            elif ext == 'txt':
                content = file_stream.read().decode('utf-8')
                translated_text = safe_translate(content)
                translated_content = translated_text.encode('utf-8')
                progress_dict[task_id] = 100
                progress_dict[f"{task_id}_result"] = (translated_content, ext)
                logging.info(f"TXT перевод успешен для task_id={task_id}")
            else:
                progress_dict[f"{task_id}_error"] = "Неподдерживаемый тип файла"
                logging.error(f"Неподдерживаемый тип файла для task_id={task_id}")
        else:
            progress_dict[task_id] = 0
            progress_dict[f"{task_id}_error"] = "Перевод документов моделью пока недоступен"
            logging.error(f"Перевод моделью недоступен для task_id={task_id}")
    except Exception as e:
        progress_dict[task_id] = 100
        progress_dict[f"{task_id}_error"] = str(e)
        logging.error(f"Исключение в translate_file_background для task_id={task_id}: {str(e)}")

@app.post("/translate-file")
async def translate_file(
    token: str,
    file: UploadFile = File(...),
    method: Optional[str] = "model",
    background_tasks: BackgroundTasks = None,
    db: Session = Depends(get_db)
):
    user = db.query(User).filter(User.token == token).first()
    if not user:
        raise HTTPException(status_code=401, detail="Неверный токен")
    if user.attempts_left <= 0:
        raise HTTPException(status_code=403, detail="Попытки закончились")

    task_id = secrets.token_hex(8)
    progress_dict[task_id] = 0

    file_bytes = await file.read()
    background_tasks.add_task(translate_file_background, task_id, file_bytes, file.filename, method)

    return {"task_id": task_id}

@app.get("/translate-file-progress")
async def translate_file_progress(task_id: str):
    def event_generator():
        last_progress = -1
        while True:
            progress = progress_dict.get(task_id, 0)
            if progress != last_progress:
                yield f"data: {progress}\n\n"
                last_progress = progress
            if progress >= 100 or f"{task_id}_error" in progress_dict:
                break
            time.sleep(0.5)
        if f"{task_id}_error" in progress_dict:
            yield f"data: error:{progress_dict[f'{task_id}_error']}\n\n"
        else:
            yield f"data: done\n\n"
    return StreamingResponse(event_generator(), media_type="text/event-stream")

@app.post("/refill-attempts", response_model=RefillResponse)
def refill_attempts(
    request: RefillRequest,
    db: Session = Depends(get_db)
):
    if request.admin_token != ADMIN_TOKEN:
        raise HTTPException(status_code=401, detail="Неверный токен администратора")
    if request.attempts <= 0:
        raise HTTPException(status_code=400, detail="Попытки должны быть положительными")
    
    user = db.query(User).filter(User.token == request.user_token).first()
    if not user:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    
    user.attempts_left += request.attempts
    try:
        db.commit()
        db.refresh(user)
        return {
            "message": "Попытки успешно пополнены",
            "new_attempts": user.attempts_left
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download-translated-file")
async def download_translated_file(task_id: str = Query(...)):
    result_key = f"{task_id}_result"
    if result_key not in progress_dict:
        raise HTTPException(status_code=404, detail="Переведенный файл не найден или не готов")
    file_bytes, ext = progress_dict[result_key]
    if ext == 'pdf':
        media_type = 'application/pdf'
    elif ext == 'docx':
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif ext == 'txt':
        media_type = 'text/plain'
    else:
        media_type = 'application/octet-stream'
    filename = f"translated_{task_id}.{ext}"
    headers = {
        "Content-Disposition": f"attachment; filename=\"{filename}\""
    }
    file_stream = io.BytesIO(file_bytes)
    return StreamingResponse(file_stream, media_type=media_type, headers=headers)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("api_combined:app", host="0.0.0.0", port=8000, reload=True)
